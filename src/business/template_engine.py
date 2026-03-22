#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
模板引擎
"""

from docxtpl import DocxTemplate
from docx import Document
from src.data.template_manager import TemplateManager
from src.business.data_manager import DataManager
import os
import re
from datetime import datetime
from pathlib import Path


class TemplateEngine:
    """模板引擎类"""
    
    def __init__(self):
        self.template_manager = TemplateManager()
        self.data_manager = DataManager()
        self._admin_fields_cache, self._member_fields_cache, self._template_fields_cache = self.data_manager.get_fields(src='template')
    
    # ======================== 获取模板元数据 =========================
    
    def get_templates(self, template_id=None):
        return self.template_manager.load_templates(template_id) 

    # ======================== 一些字段解析和格式化的辅助方法 =========================

    def _normalize_text(self, text: str) -> str:
        """归一化文本用于模糊匹配"""
        if not text:
            return ""
        return re.sub(r"[\s\-_/（）()，,。.:：]", "", str(text)).lower()

    def _build_member_field_alias_map(self) -> dict[str, str]:
        """构建成员基础字段别名映射：alias -> canonical_key"""
        alias_map: dict[str, str] = {}

        for field_def in self._member_fields_cache:
            key = field_def.get("key", "")
            if not key:
                continue

            aliases = {key}
            normalized_key = self._normalize_text(key)
            if normalized_key:
                aliases.add(normalized_key)

            # 从 fields_definition.json 中读取显式别名配置
            configured_aliases = field_def.get("aliases", []) or []
            for alias in configured_aliases:
                if alias:
                    aliases.add(str(alias))

            # 日期/时间类字段的常见中文写法兼容，例如“出生年月”映射到“出生日期”
            if "日期" in key:
                aliases.add(key.replace("日期", "年月"))
                aliases.add(key.replace("日期", "时间"))
                aliases.add(key.replace("日期", "日期时间"))
            if "时间" in key:
                aliases.add(key.replace("时间", "日期"))
                aliases.add(key.replace("时间", "年月"))

            for alias in aliases:
                alias_map[self._normalize_text(alias)] = key

        return alias_map

    def _resolve_member_basic_key(self, placeholder: str) -> str | None:
        """将占位符解析为成员基础字段 key（支持别名）"""
        normalized_placeholder = self._normalize_text(placeholder)
        if not normalized_placeholder:
            return None

        alias_map = self._build_member_field_alias_map()
        if normalized_placeholder in alias_map:
            return alias_map[normalized_placeholder]

        # 兜底：包含关系匹配（避免过度放宽，仅在长度差较小时生效）
        for alias, key in alias_map.items():
            if not alias:
                continue
            if (alias in normalized_placeholder or normalized_placeholder in alias) and abs(len(alias) - len(normalized_placeholder)) <= 2:
                return key

        return None

    def _infer_output_format(self, placeholder: str, field_key: str) -> str | None:
        """根据占位符语义推断输出格式"""
        if "年月" in placeholder and ("日期" in field_key or "时间" in field_key):
            return "YYYY年MM月"
        return None

    def _format_value(self, value, output_format: str | None):
        """按照目标格式格式化值（当前用于日期裁剪）"""
        if value is None:
            return ""
        text = str(value).strip()
        if not text or not output_format:
            return text

        if output_format == "YYYY年MM月":
            dt = None
            for fmt in ["%Y年%m月%d日", "%Y年%m月", "%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d"]:
                try:
                    dt = datetime.strptime(text, fmt)
                    break
                except Exception:
                    continue

            if dt:
                return dt.strftime("%Y年%m月")

            match = re.search(r"(\d{4})\D*(\d{1,2})", text)
            if match:
                year = match.group(1)
                month = int(match.group(2))
                return f"{year}年{month:02d}月"

        return text

    # ======================== 模板占位符解析与映射 =========================

    def get_placeholders(self, template_id: str) -> set[str]:
        """解析并返回模板中出现的所有占位符变量名（不含花括号）"""
        template_path = self.template_manager.get_template_file_path(template_id)

        if not template_path.exists():
            return set()

        doc = Document(str(template_path))

        placeholders: set[str] = set()

        # 段落中的占位符
        for paragraph in doc.paragraphs:
            for match in re.findall(r"{{(.*?)}}", paragraph.text):
                placeholders.add(match.strip())

        # 表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in re.findall(r"{{(.*?)}}", cell.text):
                        placeholders.add(match.strip())

        return placeholders
    
    def match_placehoder_def(self, placeholder: str, ) -> dict:
        for field_def in self._template_fields_cache:
            if field_def.get("is_default"):
                continue
            keywords = field_def.get("match_keywords", [])
            for keyword in keywords:
                if keyword in placeholder:
                    return {
                        "key": placeholder,
                        "type": field_def.get("type", "text"),
                        "required": field_def.get("required", False),
                        "format": field_def.get("format"),       # 这里应该先行考虑下“年月”时的处理方式
                        "display": field_def.get("display", {}),
                    }

        # 如果没有匹配，返回默认定义
        default_def = next((f for f in self._template_fields_cache if f.get("is_default")), None)
        if default_def:
            return {
                "key": placeholder,
                "type": default_def.get("type", "text"),
                "required": default_def.get("required", False),
                "display": default_def.get("display", {}),
            }

        return {
            "key": placeholder,
            "type": "text",
            "required": False,
            "display": {"order": 999},
        }
    
    def auto_map_placeholders(self, template_id: str) -> dict:
        """
        自动映射占位符到数据源
        返回格式：{
            "{{姓名}}": {
                "source": "basic_info",
                "field": "姓名",
                "auto_mapped": True
            },
            ...
        }
        """
        placeholders = self.get_placeholders(template_id)
        mapping = {}
        
        # 加载已知字段定义
        basic_fields, admin_fields = self._member_fields_cache, self._admin_fields_cache
        basic_keys = {f.get("key") for f in basic_fields}
        
        # 构建管理员字段映射表（key -> (group, key)）
        admin_key_to_group_key = {}
        for admin_field in admin_fields:
            key = admin_field.get("key", "")
            group = admin_field.get("group", "")
            if key:
                admin_key_to_group_key[key] = (group, key)
        
        # 智能匹配
        for placeholder in placeholders:
            # 1. 尝试匹配基础信息字段
            matched_basic_key = placeholder if placeholder in basic_keys else self._resolve_member_basic_key(placeholder)
            if matched_basic_key:
                mapping[f"{{{{{placeholder}}}}}"] = {
                    "source": "basic_info",
                    "field": matched_basic_key,
                    "output_format": self._infer_output_format(placeholder, matched_basic_key),
                    "auto_mapped": True
                }
                continue
            
            # 2. 尝试匹配管理员配置字段
            if placeholder in admin_key_to_group_key:
                group, key = admin_key_to_group_key[placeholder]
                mapping[f"{{{{{placeholder}}}}}"] = {
                    "source": "admin_config",
                    "group": group,
                    "key": key,
                    "auto_mapped": True
                }
                continue
            
            # 3. 默认作为模板特有字段
            mapping[f"{{{{{placeholder}}}}}"] = {
                "source": "template_data",
                "template_id": template_id,
                "field": placeholder,
                "auto_mapped": True
            }
        
        return mapping
    
    # ======================== 合并有关数据 =========================

    def merge_data_for_template(self, template_id):
        """合并数据用于模板生成（方案C：混合模式）"""
        merged_data = {}
        
        # 1. 加载管理员配置
        admin_config = self.data_manager.get_admin_config()
        
        # 2. 加载成员数据
        member_info = self.data_manager.get_member_info()
        
        # 3. 获取字段映射
        field_mapping = self.auto_map_placeholders(template_id)
        
        # 4. 根据映射合并数据
        #    已配置映射的占位符优先按照映射规则取值
        for placeholder, mapping in field_mapping.items():
            key = placeholder.strip('{}')
            value = self._get_value_by_mapping(mapping, admin_config, member_info)
            merged_data[key] = value

        # 5. 获取管理员配置的模板字段
        admin_template_data = admin_config.get("template_data", {}).get(template_id, {})
        
        # 6. 注入模板数据中所有未映射的字段（应用方案C混合模式）
        tpl_data = member_info.get('template_data', {}).get(template_id, {})
        for k, v in tpl_data.items():
            if k not in merged_data:
                # 检查管理员是否配置并锁定了该字段
                admin_field_config = admin_template_data.get(k, {})
                if isinstance(admin_field_config, dict):
                    is_locked = admin_field_config.get("locked", False)
                    admin_value = admin_field_config.get("value", "")
                else:
                    # 兼容旧格式（直接存储值）
                    is_locked = False
                    admin_value = admin_field_config if admin_field_config else ""
                
                if is_locked and admin_value:
                    # 字段被锁定，使用管理员配置的值
                    merged_data[k] = admin_value
                else:
                    # 字段未锁定，优先使用成员数据
                    merged_data[k] = v if v else admin_value
        
        # 7. 补充管理员配置但成员未填写的字段
        for k, admin_field_config in admin_template_data.items():
            if k not in merged_data:
                if isinstance(admin_field_config, dict):
                    merged_data[k] = admin_field_config.get("value", "")
                else:
                    merged_data[k] = admin_field_config if admin_field_config else ""
        
        return merged_data
    
    def _get_value_by_mapping(self, mapping, admin_config, member_info):
        """根据映射获取值"""
        source = mapping.get('source')
        
        if source == 'basic_info':
            field = mapping.get('field')
            value = member_info.get('basic_data', {}).get(field, '')
            return self._format_value(value, mapping.get('output_format'))
        
        elif source == 'admin_config':
            # 使用 group + key 从管理员配置中获取值
            group = mapping.get('group', '')
            key = mapping.get('key', '')
            return admin_config.get("basic_data", {}).get(group, {}).get(key, '')
        
        elif source == 'template_data':
            template_id = mapping.get('template_id')
            field = mapping.get('field')
            return member_info.get('template_data', {}).get(template_id, {}).get(field, '')
        
        return ''
    
    # ======================== 生成模板 =========================
    
    def generate_document(self, template_id):
        """生成 Word 文档"""
        # 1. 获取模板文件路径
        template_path = self.template_manager.get_template_file_path(template_id)

        if not template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        # 导出文件名和导出路径
        template_name = self.get_templates(template_id).get("name", "文档")
        name = self.data_manager.get_member_info("basic_data", "姓名") or "未命名"
        date_str = datetime.now().strftime("%Y%m%d")
        filename = f"{template_name}_{name}_{date_str}.docx"

        export_path = self.data_manager.get_system_settings("export_path") or "./exports"
        export_dir = Path(export_path)
        export_dir.mkdir(parents=True, exist_ok=True)

        output_path = str(export_dir / filename)
        
        # 2. 合并数据
        data = self.merge_data_for_template(template_id)
        
        # 3. 生成文档
        return self._generate_with_docxtpl(template_path, data, output_path)
    
    def _generate_with_docxtpl(self, template_path, data, output_path):
        """使用 docxtpl 生成文档"""
        doc = DocxTemplate(str(template_path))
        doc.render(data)

        # 确保输出目录存在
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    

