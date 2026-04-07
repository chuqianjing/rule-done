#!/usr/bin/env python
# -*- coding: utf-8 -*-
# Copyright (c) 2026 楚乾靖(Chu Qianjing)
# Licensed under the GNU General Public License v3.0 (GPL-3.0).
"""
模板引擎模块

本模块负责 Word 模板的占位符解析、占位符与业务数据的映射、
模板渲染和文档导出。

主要功能：
1. tpl_metadata: 读取模板元数据与模板文件
2. tpl_placeholders: 提取模板中的占位符变量
3. placeholder~datasrc: 将占位符映射到管理员/成员数据源
4. data_merge: 合并最终渲染数据
5. generate_document: 通过 docxtpl 生成导出文档

Author: 楚乾靖 (Chu Qianjing)
Date: 2026-03
"""

from datetime import datetime
from pathlib import Path
import os
import re
from docx import Document
from docxtpl import DocxTemplate
from src.application.data_manager import DataManager
from src.persistence.template_manager import TemplateManager


class TemplateEngine:
    """模板引擎类
    
    用于连接模板定义、业务数据与导出流程，提供从占位符解析到文档生成的完整能力。
    
    Attributes:
        template_manager (TemplateManager): 模板管理器，负责模板元数据与路径读取。
        data_manager (DataManager): 数据管理器，负责获取配置、成员数据与系统设置。
        _admin_fields (list[dict]): 管理员字段定义。
        _member_fields (list[dict]): 成员字段定义。
        _template_fields (list[dict]): 模板专有字段定义。
    """
    
    def __init__(self):
        """初始化模板引擎并加载字段定义缓存。"""
        self.template_manager = TemplateManager()
        self.data_manager = DataManager()
        self._admin_fields, self._member_fields, self._template_fields = self.data_manager.get_fields(src='template')
    
    # ======================== 获取模板元数据 =========================
    
    def get_templates(self, template_id=None):
        """获取模板元数据。

        Args:
            template_id (str | None): 模板ID。为 None 时返回全部模板。

        Returns:
            dict | list: TemplateManager 返回的模板数据结构。
        """
        return self.template_manager.load_templates(template_id)

    # ======================== 模板占位符解析与映射 =========================

    def get_placeholders(self, template_id: str) -> set[str]:
        """解析并返回模板中的占位符变量名（不含花括号）。

        会扫描段落与表格中的文本，识别形如 `{{变量名}}` 的占位符。

        Args:
            template_id (str): 模板ID。

        Returns:
            set[str]: 去重后的占位符名称集合。
        """
        template_path = self.template_manager.get_template_file_path(template_id)

        if not template_path.exists():
            return set()

        doc = Document(str(template_path))

        placeholders: set[str] = set()

        # 段落中的占位符
        for paragraph in doc.paragraphs:
            for match in re.findall(r"{{(.*?)}}", paragraph.text):
                placeholders.add(match.strip())

        # 表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for match in re.findall(r"{{(.*?)}}", cell.text):
                        placeholders.add(match.strip())

        return placeholders
    
    def match_placehoder_def(self, placeholder: str, ) -> dict:
        """根据占位符内容匹配字段定义。

        先按 `match_keywords` 从模板字段定义中匹配，若无匹配则回退到
        `is_default` 的默认定义；若默认定义也不存在，则返回内置兜底定义。

        Args:
            placeholder (str): 占位符名称（不含花括号）。

        Returns:
            dict: 字段定义映射，包含 key/type/required/display 等信息。
        """
        for field_def in self._template_fields:
            if field_def.get("is_default"):
                continue
            keywords = field_def.get("match_keywords", [])
            for keyword in keywords:
                if keyword in placeholder:
                    return {
                        "key": placeholder,
                        "type": field_def.get("type", "text"),
                        "required": field_def.get("required", False),
                        "format": field_def.get("format"),
                        "display": field_def.get("display", {}),
                    }

        # 如果没有匹配，返回默认定义
        default_def = next((f for f in self._template_fields if f.get("is_default")), None)
        if default_def:
            return {
                "key": placeholder,
                "type": default_def.get("type", "text"),
                "required": default_def.get("required", False),
                "display": default_def.get("display", {}),
            }

        return {
            "key": placeholder,
            "type": "text",
            "required": False,
            "display": {"order": 999},
        }
    
    def map_placeholders_to_data(self, template_id: str) -> dict:
        """将模板占位符自动映射到具体数据源。用于成员模板页的专有项字段构建、数据加载、文件生成
        映射逻辑仅发挥作用于member_template_page，admin_template_page的专有项直接显示admin_config自己的数据。

        映射优先级概述：
        1. 优先匹配成员基础字段。
        2. 其次匹配管理员基础字段。
        3. 对模板专有字段，依据管理员模板配置与成员模板数据状态决定来源。

        Args:
            template_id (str): 模板ID。

        Returns:
            dict: 占位符到数据源描述的映射。

        映射示例：
            {
                "姓名": {
                    "source": "member_basic_data",
                    "key": "姓名",
                    "order": 1,
                }
            }
        """
        placeholders = self.get_placeholders(template_id)
        mapping = {}

        member_template_data = self.data_manager.get_member_info("template_data", template_id) or {}
        admin_template_data = self.data_manager.get_admin_config("template_data", template_id) or {}

        # 先检查是否已经锁定材料
        is_locked_document = member_template_data.get("locked", False)
        if is_locked_document:
            # 若已锁定，则所有专有项均以成员模板数据为准（无论是否有值，且不显示提示）
            for placeholder in placeholders:
                mapping[placeholder] = {
                    "source": "member_template_data",
                    "is_locked": True,
                }
            return mapping

        member_keys = {f.get("key") for f in self._member_fields}
        admin_keys = {}     # 构建管理员字段键（key -> (group, key)）
        for admin_field in self._admin_fields:
            key = admin_field.get("key", "")
            group = admin_field.get("group", "")
            if key:
                admin_keys[key] = (group, key)

        
        member_template_version = member_template_data.get("version")
        admin_template_version = self.data_manager.get_admin_config("version")
        subject_to_member_template = False     # True表示以member_template_data为准，False表示以目前的数据映射逻辑处理方式为准
        # 根据两个version的时间大小关系决定是否以member_template_data为准
        if member_template_version and admin_template_version:
            member_version_time = datetime.strptime(member_template_version, "%Y.%m.%d")
            admin_version_time = datetime.strptime(admin_template_version, "%Y.%m.%d")
            if (admin_version_time - member_version_time).days >= 30:
                subject_to_member_template = True



        # 匹配
        for placeholder in placeholders:
            '''
            先成员再管理员的逻辑可实现：当成员和管理员的basic_data都有以当前placeholder为键的项，则优先显示成员的
            '''
            # 1、尝试与成员的基本字段来匹配
            if placeholder in member_keys:
                mapping[placeholder] = {
                    "source": "member_basic_data",
                    "key": placeholder,
                    "order": next((f.get("display", {}).get("order", 999) for f in self._member_fields if f.get("key") == placeholder), 999),
                }
                continue
            if placeholder == "出生年月":
                mapping[placeholder] = {
                    "source": "member_basic_data",
                    "key": "出生日期",     # 因此，若 placeholder!=key's value，此处应为 key's value
                    "order": next((f.get("display", {}).get("order", 999) for f in self._member_fields if f.get("key") == "出生日期"), 999),
                    "format": "YYYY年MM月",
                }
                continue
            # 2、尝试与管理员的基本字段来匹配
            if placeholder in admin_keys:
                group, key = admin_keys[placeholder]
                mapping[placeholder] = {
                    "source": "admin_basic_data",
                    "group": group,
                    "key": key,
                }
                continue

            '''
            模板特有字段的数据源
            '''
            # 模板特有占位符
            if subject_to_member_template:
                mapping[placeholder] = {
                    "source": "member_template_data",
                }
                continue


            member_value = member_template_data.get(placeholder, "")

            admin_field_config = admin_template_data.get(placeholder, {})
            admin_value = admin_field_config.get("value", "")
            is_locked = admin_field_config.get("locked", False)

            if is_locked:
                mapping[placeholder] = {
                    "source": "admin_template_data",
                    "is_tip": False,
                }
            elif admin_value and not member_value:
                 mapping[placeholder] = {
                    "source": "admin_template_data",
                    "is_tip": True,
                }
            else:
                mapping[placeholder] = {
                    "source": "member_template_data",
                }

        return mapping
    
    # ======================== 合并有关数据 =========================

    def merge_data_for_template(self, template_id):
        """合并模板渲染所需数据

        根据 `map_placeholders_to_data` 的映射结果，从管理员配置和成员数据中
        提取最终值并组装为渲染字典。

        Args:
            template_id (str): 模板ID。

        Returns:
            dict: 可直接传入 docxtpl 的渲染数据。

        Notes:
            - 当字段格式为 `YYYY年MM月` 时，会基于 `出生日期` 做年月转换。
            - 默认日期值 `1000年1月1日` 会被替换为 `无`。
        """
        merged_data = {}

        admin_config = self.data_manager.get_admin_config()
        member_info = self.data_manager.get_member_info()
        
        placeholder_mapping = self.map_placeholders_to_data(template_id)

        for placeholder, mapping in placeholder_mapping.items():

            data_src = mapping.get('source')
            group = mapping.get('group', '')
            key = mapping.get('key', '')
            format = mapping.get('format', '')

            if data_src == "member_basic_data":
                value = member_info.get('basic_data', {}).get(key, '')
                if format == "YYYY年MM月" and value:
                    dt = datetime.strptime(value, "%Y年%m月%d日")
                    value = f"{dt.year}年{dt.month}月"
                if value == "1000年1月1日":   # 处理默认日期值
                    value = "无"
            elif data_src == 'admin_basic_data':
                value = admin_config.get("basic_data", {}).get(group, {}).get(key, '')
            elif data_src == 'admin_template_data':
                value = admin_config.get('template_data', {}).get(template_id, {}).get(placeholder, {}).get("value", '')
            elif data_src == 'member_template_data':
                if mapping.get("is_locked", False):
                    tpl_data = member_info.get('template_data', {}).get(template_id, {})
                    placeholder_dict = {}
                    for k, v in tpl_data.get("basic_entry", {}).items():
                        placeholder_dict[k] = v
                    for k, v in tpl_data.get("template_entry", {}).items():
                        placeholder_dict[k] = v
                    value = placeholder_dict.get(placeholder, '')
                else:
                    value = member_info.get('template_data', {}).get(template_id, {}).get(placeholder, '')
            else:
                value = ''
            merged_data[placeholder] = value

        return merged_data
    
    # ======================== 生成模板 =========================
    
    def generate_document(self, template_id):
        """生成并导出 Word 文档

        流程包括：获取模板文件、组装导出路径、合并数据、调用渲染器生成文档。

        Args:
            template_id (str): 模板ID。

        Returns:
            str: 生成后的文档完整路径。
        """
        # 1. 获取模板文件路径
        template_path = self.template_manager.get_template_file_path(template_id)

        if not template_path.exists():
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        # 导出文件名和导出路径
        template_name = self.get_templates(template_id).get("name", "文档")
        name = self.data_manager.get_member_info("basic_data", "姓名") or "未命名"
        date_str = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"{template_name}_{name}_{date_str}.docx"

        export_path = self.data_manager.get_system_settings("export_path") or "./exports"
        export_dir = Path(export_path)
        export_dir.mkdir(parents=True, exist_ok=True)

        output_path = str(export_dir / filename)
        
        # 2. 合并数据
        data = self.merge_data_for_template(template_id)
        
        # 3. 生成文档
        return self._generate_with_docxtpl(template_path, data, output_path)
    
    def _generate_with_docxtpl(self, template_path, data, output_path):
        """使用 docxtpl 渲染并保存文档

        Args:
            template_path (str | Path): 模板文件路径。
            data (dict): 模板渲染数据。
            output_path (str): 输出文件路径。

        Returns:
            str: 输出文件路径。
        """
        doc = DocxTemplate(str(template_path))
        doc.render(data)

        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)
        
        return output_path
    

